from pymongo import MongoClient
from bson.objectid import ObjectId
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.shared import RGBColor
#------------------------------------資料取得----------------------------------------
def mongo_connect():
    try:
        client = MongoClient("address")
        return client["pv"]
    except Exception as e:
        print(e)

#輸入id，輸出plant_id, solar_ID, meter_ID, pr_ID
def id_identify(pv, id):
    plant = pv["plant"].find_one(
        {
            "_id": ObjectId(id)
        }
    )
    equipment = pv["equipment"].find_one(
        {
            "_id": ObjectId(id)
        }
    )
    if plant and not equipment:
        plant_id = id
        name = plant["name"]
        solar_ID = solar_meter_ID(pv, name)
        meter_ID = id
        pr_ID = id
    elif not plant and equipment:
        name = equipment["PV"]
        plant = pv["plant"].find_one(
            {
                "name": name
            }
        )
        plant_id = str(plant["_id"])
        PV_str = equipment["PV"]
        pv_lgroup = equipment["lgroup"]
        group = equipment["name"]
        solar_ID = solar_meter_ID(pv, PV_str, pv_lgroup, group)
        meter_ID = id
        pr_ID = id
    return plant_id, solar_ID, meter_ID, pr_ID

#案場資訊
def field_imformation(pv, object_id):
    plant = pv["plant"].find_one(
        {
            "_id": ObjectId(object_id)
        }
    )

    name = plant["name"]
    field_position = []
    field_position.append(plant["plant_address"])
    field_position.append(plant["coordinates"])
    capacity = plant["capacity"]
    return name, field_position, capacity

#日照計ID
def solar_meter_ID(pv, PV, lgroup=None, group=None):
    equipment = pv["equipment"]
    if not (lgroup and group):
        meter = equipment.find_one(
            {
                "PV": PV,
                "type": "sun",
                "main_sun": 1
            }
        )
    elif lgroup and not group:
        meter = equipment.find_one(
            {
                "PV": PV,
                "lgroup": lgroup,
                "type": "sun",
                "main_sun": 1
            }
        )
    elif lgroup and group:
        meter = equipment.find_one(
            {
                "PV": PV,
                "lgroup": lgroup,
                "group": group,
                "type": "sun",
                "main_sun": 1
            }
        )
    ID = str(meter["_id"])
    return ID

#PR跟meterID
def PR_and_meter_ID(pv, PV):
    equipment = pv["equipment"]
    meter = equipment.find_one(
        {
            "PV": PV,
            "type": "pv_group"
        }
    )
    ID = str(meter["_id"])
    return ID

#日照計的data list
def irrh_cal(pv, ID, times, time_interval, round_number=1):
    irrh_data = []
    irrh_cal = pv["irrh_cal"]
    for time in times:
        irrh = irrh_cal.find_one(
            {
                "ID": ID,
                "time": time, 
                "time_interval": time_interval
            }
        )
        if irrh:
            if irrh["irrh"]:
                if type(round_number) == int:
                    irrh_data.append(round(irrh["irrh"], round_number))
                else:
                    round_number = round(round_number)
                    irrh_data.append(round(irrh["irrh"], round_number))
            else:
                irrh_data.append("---")
        else:
            irrh_data.append("---")
    return irrh_data

#meter的data list
def meter_cal(pv, ID, times, time_interval, round_number=1):
    meter_data = []
    meter_cal = pv["meter_cal"]
    for time in times:
        meter = meter_cal.find_one(
            {
                "ID": ID,
                "time": time, 
                "time_interval": time_interval
            }
        )
        if meter:
            if meter["kwh"]:
                if type(round_number) == int:
                    meter_data.append(round(meter["kwh"], round_number))
                else:
                    round_number = round(round_number)
                    meter_data.append(round(meter["kwh"], round_number))
            else:
                meter_data.append("---")
        else:
            meter_data.append("---")
    return meter_data

#pr的data list
def pr_cal(pv, ID, times, time_interval, round_number=1):
    pr_data = []
    pr_cal = pv["pr_cal"]
    for time in times:
        pr = pr_cal.find_one(
            {
                "ID": ID,
                "time": time, 
                "time_interval": time_interval
            }
        )
        if pr:
            if pr["pr"]:
                if type(round_number) == int:
                    pr_data.append(round(pr["pr"], round_number))
                else:
                    round_number = round(round_number)
                    pr_data.append(round(pr["pr"], round_number))
            else:
                pr_data.append("---")
        else:
            pr_data.append("---")
    return pr_data

#判斷是否為閏年並返回天數
def leap_year(year):
    if (year % 4) == 0:
        if (year % 100) == 0:
            if (year % 400) == 0:
                month_day = 29
            else:
                month_day = 28
        else:
            month_day = 29
    else:
        month_day = 28
    return month_day

#所有需要的時間
def set_time_interval(start_time, end_time, time_interval):
    during = []
    large_month = [1, 3, 5, 7, 8, 10, 12]
    small_month = [4, 6, 9, 11]
    start = str(start_time).split()
    start_date = start[0]
    start_hour_string = start[1]

    end = str(end_time).split()
    end_date = end[0]
    end_hour_string = end[1]

    if time_interval == "year":
        start_year = start_date.split("-")[0]
        end_year = end_date.split("-")[0]
        component = "-01-01 00:00:00"
        date = start_year + component
        during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
        while start_year < end_year:
            start_year = str(int(start_year) + 1)
            date = start_year + component
            during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
    elif time_interval == "month":
        start_list = start_date.split("-")
        start_year = start_list[0]
        start_month = int(start_list[1])
        end_list = end_date.split("-")
        end_year = end_list[0]
        end_month = int(end_list[1])
        component = "-01 00:00:00"
        date = start_year + "-" + str(start_month) + component
        during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
        while start_year < end_year or start_month < end_month:
            start_month += 1
            if start_month == 13:
                start_year = str(int(start_year) + 1)
                start_month = 1
            date = start_year + "-" + str(start_month) + component
            during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
    elif time_interval == "day":
        start_list = start_date.split("-")
        start_year = int(start_list[0])
        start_month = int(start_list[1])
        start_day = int(start_list[2])
        end_list = end_date.split("-")
        end_year = int(end_list[0])
        end_month = int(end_list[1])
        end_day = int(end_list[2])
        component = " 00:00:00"
        date = str(start_year) + "-" + str(start_month) + "-" + str(start_day) + component
        during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
        while start_year < end_year or start_month < end_month or start_day < end_day:
            start_day += 1
            if start_month in large_month:
                if start_day == 32:
                    start_day = 1
                    start_month += 1
            elif start_month in small_month:
                if start_day == 31:
                    start_day = 1
                    start_month += 1
            elif start_month == 2:
                leap_month_day = leap_year(start_year)
                if start_day == leap_month_day + 1:
                    start_day = 1
                    start_month += 1
            if start_month == 13:
                start_month = 1
                start_year += 1
            date = str(start_year) + "-" + str(start_month) + "-" + str(start_day) + component
            during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
    elif time_interval == "hour":
        start_list = start_date.split("-")
        start_year = int(start_list[0])
        start_month = int(start_list[1])
        start_day = int(start_list[2])
        start_hour = int(start_hour_string.split(":")[0])
        end_list = end_date.split("-")
        end_year = int(end_list[0])
        end_month = int(end_list[1])
        end_day = int(end_list[2])
        end_hour = int(end_hour_string.split(":")[0])
        component = ":00:00"
        date = str(start_year) + "-" + str(start_month) + "-" + str(start_day) + " " + str(start_hour) + component
        during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
        while start_year < end_year or start_month < end_month or start_day < end_day or start_hour < end_hour:
            start_hour += 1
            if start_hour == 24:
                start_hour = 0
                start_day += 1
                if start_month in large_month:
                    if start_day == 32:
                        start_day = 1
                        start_month += 1
                elif start_month in small_month:
                    if start_day == 31:
                        start_day = 1
                        start_month += 1
                elif start_month == 2:
                    leap_month_day = leap_year(start_year)
                    if start_day == leap_month_day + 1:
                        start_day = 1
                        start_month += 1
                if start_month == 13:
                    start_month = 1
                    start_year += 1
            date = str(start_year) + "-" + str(start_month) + "-" + str(start_day) + " " + str(start_hour) + component
            during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
    elif time_interval == "15min":
        start_list = start_date.split("-")
        start_year = int(start_list[0])
        start_month = int(start_list[1])
        start_day = int(start_list[2])
        start_hour_list = start_hour_string.split(":")
        start_hour = int(start_hour_list[0])
        start_minute = int(start_hour_list[1])
        end_list = end_date.split("-")
        end_year = int(end_list[0])
        end_month = int(end_list[1])
        end_day = int(end_list[2])
        end_hour_list = end_hour_string.split(":")
        end_hour = int(end_hour_list[0])
        end_minute = int(end_hour_list[1])
        component = ":00"
        date = str(start_year) + "-" + str(start_month) + "-" + str(start_day) + " " + str(start_hour) + ":" + str(start_minute) + component
        during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
        while start_year < end_year or start_month < end_month or start_day < end_day or start_hour < end_hour or start_minute < end_minute:
            start_minute += 15
            if start_minute == 60:
                start_hour += 1
                start_minute = 0
                if start_hour == 24:
                    start_hour = 0
                    start_day += 1
                    if start_month in large_month:
                        if start_day == 32:
                            start_day = 1
                            start_month += 1
                    elif start_month in small_month:
                        if start_day == 31:
                            start_day = 1
                            start_month += 1
                    elif start_month == 2:
                        leap_month_day = leap_year(start_year)
                        if start_day == leap_month_day + 1:
                            start_day = 1
                            start_month += 1
                    if start_month == 13:
                        start_month = 1
                        start_year += 1
            date = str(start_year) + "-" + str(start_month) + "-" + str(start_day) + " " + str(start_hour) + ":" + str(start_minute) + component
            during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))

    return during

#頁首資訊
def company_imformation(pv, object_id):
    plant = pv["plant"].find_one(
        {
            "_id": ObjectId(object_id)
        }
    )
    company_name = plant["paets_info"]["unit"]
    tel = plant["paets_info"]["TEL"]
    imformation = []
    imformation.append(company_name)
    imformation.append(tel)
    return imformation

def project_name(pv, equipment_id):
    equipment = pv["equipment"].find_one(
        {
            "_id": ObjectId(equipment_id)
        }
    )
    name = equipment["PV"] + "-" + equipment["lgroup"] + "-" + equipment["name"]
    return name

def imformation_data(project_name, date, position, capacity):
    imformation = {
        "專案名稱": project_name,
        "日期": date,
        "案場位置": position,
        "設置容量": capacity + "kW"
    }
    return imformation

#---------------------------報表製作-----------------------------------------------------

#初始設定
def document_initial(doc):
    #設定邊界
    section = doc.sections[0]
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    #設定字型
    doc.styles['Normal'].font.size = Pt(12)
    #設定字體(英文)
    doc.styles['Normal'].font.name = 'Times New Roman'
    #中文
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

#資料表格
def add_table(doc, head_datas, datas, width, style=None):
    table = doc.add_table(rows=1, cols=len(head_datas), style=("Table Grid" if style is None else style))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    head_cells = table.rows[0].cells
    #表格標頭文字
    for index, head_item in enumerate(head_datas):
        head_cells[index].text = head_item
        head_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        head_cells[index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #表格資料
    for data in datas:
        row_cells = table.add_row().cells
        for index, cell in enumerate(row_cells):
            cell.text = str(data[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # #表格內寬度
    # for index, column in enumerate(table.columns):
    #     for cell in column.cells:
    #         cell.width = Cm(width[index])

    #表格內高度
    for index, row in enumerate(table.rows):
        if index == 0:
            row.height = Cm(0.6)
        else:
            row.height = Cm(0.64)

#刪除段落、圖片等
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    #p._p = p._element = None
    paragraph._p = paragraph._element = None

#資訊欄
def imformation_table(doc, imformation_data):
    width = [3.9, 15.1]
    rows = 0
    index = 0
    keys = []
    row_index = 0
    key_index = 0

    #設定標題
    heading = doc.add_paragraph("")
    heading.style = doc.styles['Heading 1']
    #段前間距
    heading.paragraph_format.space_before = Pt(0)
    #段後間距
    heading.paragraph_format.space_after = Pt(15)

    run = heading.add_run("太陽光電系統維運每週報表")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

    
    # run.space_before = Pt(9)
    # run.space_after = Pt(9)
    # run.line_spacing = 1.5

    for value in imformation_data.values():
        if type(value) == list:
            rows += len(value)
        elif type(value) == str:
            rows += 1
    
    table = doc.add_table(rows=rows, cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    #table.cell(0, 0).text = "表角"

    for key in imformation_data.keys():
        keys.append(key)

    while row_index < rows:
        if type(imformation_data[keys[key_index]]) == list:
            for content in imformation_data[keys[key_index]]:
                row = table.rows[row_index]
                row.cells[0].text = keys[key_index]
                row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row.cells[1].text = content
                row_index += 1
        else:
            row = table.rows[row_index]
            row.cells[0].text = keys[key_index]
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row.cells[1].text = imformation_data[keys[key_index]]
            row_index += 1
        key_index += 1

    for index, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = Cm(width[index])


    #合併重複名稱
    col = table.columns[0]
    row_index = 0
    while row_index < rows-1:
        if col.cells[row_index].text == col.cells[row_index+1].text:
            content = col.cells[row_index].text
            col.cells[row_index].merge(col.cells[row_index+1])
            col.cells[row_index].text = content
            col.cells[row_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            col.cells[row_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_index += 1

#設定頁首
def header(doc, imformations, logo):
    section = doc.sections[0]
    header = section.header
    # paragraph = header.paragraphs[0]
    paragraph = header.add_paragraph()
    #段落間距
    paragraph.paragraph_format.space_after = Pt(0)
    #單行間距(長度為絕對距離，浮點數為間距)
    paragraph.paragraph_format.line_spacing = 1.0
    #刪除首段
    delete_paragraph(header.paragraphs[0])
    tab_stops = paragraph.paragraph_format.tab_stops
    logo_run = paragraph.add_run()


    logo_run.add_picture(logo, height=Inches(0.65))
    text_run = paragraph.add_run('\t')
    text_run = paragraph.add_run(imformations[0])
    #設置右邊文字位置tab
    margin_end = Inches(section.page_width.inches - (section.left_margin.inches + section.right_margin.inches))
    tab_stop = tab_stops.add_tab_stop(margin_end, WD_TAB_ALIGNMENT.RIGHT)
    #字體大小、斜體
    text_run.font.size = Pt(14)
    text_run.font.italic = True

    for index, imformation in enumerate(imformations):
        if index == 0:
            #paragraph = header.paragraphs[0]
            continue
        else:
            paragraph = header.add_paragraph()
            #pass
        run = paragraph.add_run(""+imformation)
        if index == 0:
            #run.font.size = Pt(14)
            continue
        else:
            run.font.size = Pt(10)
        run.font.italic = True
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
def table_data(dates, irrh_data, meter_data, pr_data):
    data = []
    for date, irrh, meter, pr in zip(dates, irrh_data, meter_data, pr_data):
        data.append([str(date).replace("-", "/"), irrh, meter, pr])
    return data

def report(name, header_imformation, imformation, table_head_datas, table_datas):
    logo = "title_picture.jpg"

    #總長約19cm
    width = [4.71, 5.49, 4.4, 4.4]

    doc = Document()
    #初始設定
    document_initial(doc)
    #設定頁首
    header(doc, header_imformation, logo)
    #資訊欄
    imformation_table(doc, imformation)
    p = doc.add_paragraph()
    #資料表格
    add_table(doc, table_head_datas, table_datas, width=width)
    doc_name = name + ".docx"
    doc.save(doc_name)


if __name__ == "__main__":
    pv = mongo_connect()
    plant_id = "5e8d4c884a11d7e11cd2050e"
    equipment_id = "5e8d4c884a11d7e11cd20521"
    plant_id, solar_ID, meter_ID, pr_ID = id_identify(pv, equipment_id)
    name, field_position, capacity = field_imformation(pv, plant_id)

    time = "2021-11-20 06:00:00"
    time = datetime.strptime(time, '%Y-%m-%d %H:%M:%S')
    time1 = "2022-02-25 10:00:00"
    time1 = datetime.strptime(time1, '%Y-%m-%d %H:%M:%S')
    time_interval = "day"
    time_list = set_time_interval(time, time1, time_interval)

    irrh_data = irrh_cal(pv, solar_ID, time_list, time_interval)
    meter_data = meter_cal(pv, meter_ID, time_list, time_interval)
    pr_data = pr_cal(pv, pr_ID, time_list, time_interval)


    header_data = company_imformation(pv, plant_id)

    name = project_name(pv, pr_ID)
    date = str(time) + "~" + str(time1)
    imformation_dict = imformation_data(name, date, field_position, str(capacity))
    table_head_datas = ["時間", "日照量", "發電量", "PR"]
    data = table_data(time_list, irrh_data, meter_data, pr_data)
    report(name, header_data, imformation_dict, table_head_datas, data)
    