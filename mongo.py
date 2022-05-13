from pymongo import MongoClient
from bson.objectid import ObjectId
from datetime import datetime, date, timedelta
from dateutil.relativedelta import relativedelta

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.shared import RGBColor

import math
import os
#------------------------------------資料取得----------------------------------------
def mongo_connect():
    hostname = os.getenv("MONGODB_HOSTNAME", "localhost")
    port = os.getenv("MONGODB_PORT", "27017")
    username = os.getenv("MONGODB_USERNAME", "root")
    password = os.getenv("MONGODB_PASSWORD", "pc152")
    connect_string = "mongodb://{}:{}@{}:{}/".format(username, password, hostname, port)
    try:
        client = MongoClient(connect_string)
        return client["pv"]
    except Exception as e:
        print(e)

#輸入id，輸出plant_id, solar_ID, meter_ID, pr_ID
def id_identify(pv, id):
    plant = pv["plant"].find_one(
        {
            "_id": ObjectId(id)
        }
    )
    equipment = pv["equipment"].find_one(
        {
            "_id": ObjectId(id)
        }
    )
    if plant and not equipment:
        plant_id = id
        name = plant["name"]
        solar_ID = solar_meter_ID(pv, name)
        meter_ID = id
        pr_ID = id
        collection = plant.get("collection", " ")
    elif not plant and equipment:
        name = equipment["PV"]
        plant = pv["plant"].find_one(
            {
                "name": name
            }
        )
        collection = equipment.get("collection", " ")
        plant_id = str(plant["_id"])
        PV_str = equipment["PV"]
        if equipment["type"] == "pv_lgroup":
            pv_lgroup = equipment["name"]
            group = None
        elif equipment["type"] == "pv_group":
            pv_lgroup = equipment["lgroup"]
            group = equipment["name"]
        solar_ID = solar_meter_ID(pv, PV_str, pv_lgroup, group)
        meter_ID = id
        pr_ID = id
    return plant_id, solar_ID, meter_ID, pr_ID, collection

#案場資訊
def field_imformation(pv, object_id):
    plant = pv["plant"].find_one(
        {
            "_id": ObjectId(object_id)
        }
    )

    name = plant["name"]
    field_position = []
    field_position.append(plant["plant_address"])
    field_position.append(plant["coordinates"])
    capacity = plant["capacity"]
    return name, field_position, capacity

#日照計ID
def solar_meter_ID(pv, PV, lgroup=None, group=None):
    equipment = pv["equipment"]
    if not lgroup and not group:
        meter = equipment.find_one(
            {
                "PV": PV,
                "type": "sun",
                "main_sun": 1
            }
        )
    elif lgroup and not group:
        meter = equipment.find_one(
            {
                "PV": PV,
                "lgroup": lgroup,
                "type": "sun",
                "main_sun": 1
            }
        )
    elif lgroup and group:
        meter = equipment.find_one(
            {
                "PV": PV,
                "lgroup": lgroup,
                "group": group,
                "type": "sun",
                "main_sun": 1
            }
        )
    #print(meter)
    if meter:
        ID = str(meter.get("_id", ""))
    else:
        ID = ""
    return ID

#PR跟meterID
def PR_and_meter_ID(pv, PV):
    equipment = pv["equipment"]
    meter = equipment.find_one(
        {
            "PV": PV,
            "type": "pv_group"
        }
    )
    ID = str(meter["_id"])
    return ID

#日照計的data list
def irrh_cal(pv, ID, times, time_interval, round_number=1):
    irrh_data = []
    irrh_cal = pv["irrh_cal"]
    for time in times:
        if len(ID) > 0:
            irrh = irrh_cal.find_one(
                {
                    "ID": ID,
                    "time": time, 
                    "time_interval": time_interval
                }
            )
            if irrh:
                if irrh["irrh"]:
                    if type(round_number) == int:
                        irrh_data.append(round(irrh["irrh"], round_number))
                    else:
                        round_number = round(round_number)
                        irrh_data.append(round(irrh["irrh"], round_number))
                else:
                    irrh_data.append("---")
            else:
                irrh_data.append("---")
        else:
            irrh_data.append("---")
    return irrh_data

#meter的data list
def meter_cal(pv, ID, times, time_interval, round_number=1):
    meter_data = []
    meter_cal = pv["meter_cal"]
    for time in times:
        meter = meter_cal.find_one(
            {
                "ID": ID,
                "time": time, 
                "time_interval": time_interval
            }
        )
        if meter:
            if meter["kwh"]:
                if type(round_number) == int:
                    meter_data.append(round(meter["kwh"], round_number))
                else:
                    round_number = round(round_number)
                    meter_data.append(round(meter["kwh"], round_number))
            else:
                meter_data.append("---")
        else:
            meter_data.append("---")
    return meter_data

#pr的data list
def pr_cal(pv, ID, times, time_interval, round_number=1):
    pr_data = []
    pr_cal = pv["pr_cal"]
    for time in times:
        pr = pr_cal.find_one(
            {
                "ID": ID,
                "time": time, 
                "time_interval": time_interval
            }
        )
        if pr:
            if pr["pr"]:
                if type(round_number) == int:
                    pr_data.append(round(pr["pr"], round_number))
                else:
                    round_number = round(round_number)
                    pr_data.append(round(pr["pr"], round_number))
            else:
                pr_data.append("---")
        else:
            pr_data.append("---")
    return pr_data

#判斷是否為閏年並返回天數
def leap_year(year):
    if (year % 4) == 0:
        if (year % 100) == 0:
            if (year % 400) == 0:
                month_day = 29
            else:
                month_day = 28
        else:
            month_day = 29
    else:
        month_day = 28
    return month_day

#所有需要的時間
def set_time_interval(start_time, end_time, time_interval):
    during = []
    large_month = [1, 3, 5, 7, 8, 10, 12]
    small_month = [4, 6, 9, 11]
    start = str(start_time).split()
    start_date = start[0]

    end = str(end_time).split()
    end_date = end[0]

    if time_interval == "year":
        start_year = start_date.split("-")[0]
        end_year = end_date.split("-")[0]
        component = "-01-01 00:00:00"
        date = start_year + component
        during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
        while start_year < end_year:
            start_year = str(int(start_year) + 1)
            date = start_year + component
            during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
    elif time_interval == "month":
        start_list = start_date.split("-")
        start_year = start_list[0]
        start_month = int(start_list[1])
        end_list = end_date.split("-")
        end_year = end_list[0]
        end_month = int(end_list[1])
        component = "-01 00:00:00"
        date = start_year + "-" + str(start_month) + component
        during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
        while start_year < end_year or start_month < end_month:
            start_month += 1
            if start_month == 13:
                start_year = str(int(start_year) + 1)
                start_month = 1
            date = start_year + "-" + str(start_month) + component
            during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
    elif time_interval == "day":
        start_list = start_date.split("-")
        start_year = int(start_list[0])
        start_month = int(start_list[1])
        start_day = int(start_list[2])
        end_list = end_date.split("-")
        end_year = int(end_list[0])
        end_month = int(end_list[1])
        end_day = int(end_list[2])
        component = " 00:00:00"
        date = str(start_year) + "-" + str(start_month) + "-" + str(start_day) + component
        during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
        while start_year < end_year or start_month < end_month or start_day < end_day:
            start_day += 1
            if start_month in large_month:
                if start_day == 32:
                    start_day = 1
                    start_month += 1
            elif start_month in small_month:
                if start_day == 31:
                    start_day = 1
                    start_month += 1
            elif start_month == 2:
                leap_month_day = leap_year(start_year)
                if start_day == leap_month_day + 1:
                    start_day = 1
                    start_month += 1
            if start_month == 13:
                start_month = 1
                start_year += 1
            date = str(start_year) + "-" + str(start_month) + "-" + str(start_day) + component
            during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
    elif time_interval == "hour":
        start_list = start_date.split("-")
        start_year = int(start_list[0])
        start_month = int(start_list[1])
        start_day = int(start_list[2])
        start_hour = 0
        end_list = end_date.split("-")
        end_year = int(end_list[0])
        end_month = int(end_list[1])
        end_day = int(end_list[2])
        end_hour = 23
        component = ":00:00"
        date = str(start_year) + "-" + str(start_month) + "-" + str(start_day) + " " + str(start_hour) + component
        during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
        while start_year < end_year or start_month < end_month or start_day < end_day or start_hour < end_hour:
            start_hour += 1
            if start_hour == 24:
                start_hour = 0
                start_day += 1
                if start_month in large_month:
                    if start_day == 32:
                        start_day = 1
                        start_month += 1
                elif start_month in small_month:
                    if start_day == 31:
                        start_day = 1
                        start_month += 1
                elif start_month == 2:
                    leap_month_day = leap_year(start_year)
                    if start_day == leap_month_day + 1:
                        start_day = 1
                        start_month += 1
                if start_month == 13:
                    start_month = 1
                    start_year += 1
            date = str(start_year) + "-" + str(start_month) + "-" + str(start_day) + " " + str(start_hour) + component
            during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
    elif time_interval == "15min":
        start_list = start_date.split("-")
        start_year = int(start_list[0])
        start_month = int(start_list[1])
        start_day = int(start_list[2])
        start_hour = 0
        start_minute = 0
        end_list = end_date.split("-")
        end_year = int(end_list[0])
        end_month = int(end_list[1])
        end_day = int(end_list[2])
        end_hour = 23
        end_minute = 45
        component = ":00"
        date = str(start_year) + "-" + str(start_month) + "-" + str(start_day) + " " + str(start_hour) + ":" + str(start_minute) + component
        during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))
        while start_year < end_year or start_month < end_month or start_day < end_day or start_hour < end_hour or start_minute < end_minute:
            start_minute += 15
            if start_minute == 60:
                start_hour += 1
                start_minute = 0
                if start_hour == 24:
                    start_hour = 0
                    start_day += 1
                    if start_month in large_month:
                        if start_day == 32:
                            start_day = 1
                            start_month += 1
                    elif start_month in small_month:
                        if start_day == 31:
                            start_day = 1
                            start_month += 1
                    elif start_month == 2:
                        leap_month_day = leap_year(start_year)
                        if start_day == leap_month_day + 1:
                            start_day = 1
                            start_month += 1
                    if start_month == 13:
                        start_month = 1
                        start_year += 1
            date = str(start_year) + "-" + str(start_month) + "-" + str(start_day) + " " + str(start_hour) + ":" + str(start_minute) + component
            during.append(datetime.strptime(date, '%Y-%m-%d %H:%M:%S'))

    return during

#頁首資訊
def company_imformation(pv, object_id):
    plant = pv["plant"].find_one(
        {
            "_id": ObjectId(object_id)
        }
    )
    company_name = plant.get('paets_info', {}).get('unit', '')
    tel = plant.get('paets_info', {}).get('TEL', '')
    imformation = []
    imformation.append(company_name)
    imformation.append(tel)
    return imformation

def project_name(pv, id, time_interval):
    plant = pv["plant"].find_one(
        {
            "_id": ObjectId(id)
        }
    )
    equipment = pv["equipment"].find_one(
        {
            "_id": ObjectId(id)
        }
    )
    if plant and not equipment:
        name = plant["name"]
    elif not plant and equipment:
        PV_str = equipment["PV"]
        if equipment["type"] == "pv_lgroup":
            lgroup = equipment["name"]
            group = None
        elif equipment["type"] == "pv_group":
            lgroup = equipment["lgroup"]
            group = equipment["name"]

        if lgroup and group:
            name = PV_str + "-" + lgroup + "-" + group
        elif lgroup and not group:
            name = PV_str + "-" + lgroup
        else:
            name = PV_str
    
    today = date.today()
    delta = timedelta(days=1)
    yesterday = today - delta
    year = yesterday.year
    month = yesterday.month

    if time_interval == "day":
        if month < 10:
            doc_name = name + "_" + str(year) + "-" + "0" + str(month) + ".docx"
        else:
             doc_name = name + "_" + str(year) + "-" + str(month) + ".docx"
    elif time_interval == "month":
        doc_name = name + "_" + str(year) + ".docx"
    else:
        doc_name = name + "_" + str(yesterday) + "_" + time_interval + ".docx"
    return doc_name

def imformation_data(project_name, date, position, capacity, max_val, avg_val="0"):
    if max_val == "0" and avg_val == "0":
        imformation = {
            "專案名稱": project_name,
            "日期": date,
            "案場位置": position,
            "設置容量": capacity + "kW",
            "最大功率": "---",
            # "平均功率": "---"
        }
    elif max_val != "0" and avg_val == "0":
        imformation = {
            "專案名稱": project_name,
            "日期": date,
            "案場位置": position,
            "設置容量": capacity + "kW",
            "最大功率": max_val + "kW",
            # "平均功率": "---"
        }
    else:
        imformation = {
            "專案名稱": project_name,
            "日期": date,
            "案場位置": position,
            "設置容量": capacity + "kW",
            "最大功率": max_val + "kW",
            # "平均功率": avg_val + "kW"
        }
    return imformation

def avg_max_value(pv, ID, collection, start_time, end_time, round_number=1, avg_status=False):
    if len(collection) > 0:
        data_collection = pv[collection]
        end_time = str(end_time).split(" ")[0] + " 23:59:59"
        end_time = datetime.strptime(end_time, '%Y-%m-%d %H:%M:%S')
        max = 0
        max_value = data_collection.find(
            {
                "ID": ID,
                "time":{"$gte": start_time, "$lt": end_time}
            }
        ).sort("p",-1).limit(1)
        if avg_status:
            avg_data = data_collection.find(
                {
                    "ID": ID,
                    "time": {"$gte": start_time, "$lt": end_time},
                    "p": {"$nin": [None, 0]}
                }
            )
            avg_list = []
            for average in avg_data:
                if average["p"] != None:
                    avg_list.append(average["p"])
            if len(avg_list) > 0:
                avg = round(sum(avg_list)/len(avg_list), round_number)
            else:
                avg = 0
        else:
            avg = 0

        if max_value:
            for maximum in max_value:
                if maximum.get("p", 0) != None:
                    max = round(maximum.get("p", 0), round_number)
                else:
                    max = 0
    else:
        max = 0
        avg = 0
    if avg_status:
        return max, avg
    else:
        return max

#---------------------------取得警報-----------------------------------------------------
def alarm_get(db, request_dict):
    alarm_filter = {}
    alarm_filter['level'] = {'$gte': 0, '$lt': 4}
    if request_dict.get('time', {}).get('mode', '') == 'single':
        try:
            start_date = datetime.strptime(request_dict['time']['start_date'], '%Y-%m-%d')
            end_date = start_date + timedelta(days=1)
            alarm_filter['time'] = {'$gte': start_date, '$lt': end_date}
        except Exception as e:
            print(e)
    elif request_dict.get('time', {}).get('mode', '') == 'interval':
        try:
            start_date = datetime.strptime(request_dict['time']['start_date'], '%Y-%m-%d')
            end_date = datetime.strptime(request_dict['time']['end_date'], '%Y-%m-%d')
            alarm_filter['time'] = {'$gte': start_date, '$lt': end_date+timedelta(days=1)}
        except Exception as e:
            print(e)
    elif request_dict.get('time', {}).get('mode', '') in ['today', 'week', 'month', 'year']:
        today =datetime.today()
        if request_dict['time']['mode'] == 'today':
            start_date = datetime.combine(today, datetime.time.min)
            end_date = start_date + timedelta(days=1)
        elif request_dict['time']['mode'] == 'week':
            start_date=datetime.combine(today, datetime.time.min) - timedelta(days=today.weekday())
            end_date = datetime.now()
        elif request_dict['time']['mode'] == 'month':
            start_date = datetime(year=today.year, month=today.month, day=1)
            end_date = datetime.now()
        elif request_dict['time']['mode'] == 'year':
            start_date = datetime(year=today.year, month=1, day=1)
            end_date = datetime.now()
        alarm_filter['time'] = {'$gte': start_date, '$lt': end_date}
    equip_filter = {}
    if request_dict.get('equip_type', '') != 'all':
        try:
            if request_dict['equip_type'] != 'iot':
                equip_filter['type'] = request_dict['equip_type']
            else:
                equip_filter['true_type'] = request_dict['equip_type']
        except Exception as e:
            print(e)
    alarm_filter['ID'] = {'$in': []}
    # plant buffer. 警報對應案場資料
    plant_buffer = {}
    for plant in db.plant.find({}):
        plant_buffer[plant.get('name', '')] = plant
    equipment_buffer = {}
    iot_buffer = {}
    if request_dict.get('plant', {}).get('all', True) == False:
        try:
            ID_list = request_dict['plant']['ID']
            col_list = request_dict['plant']['col']
            for i, _ID in enumerate(ID_list):
                try:
                    if request_dict['equip_type'] == 'all': 
                        alarm_filter['ID']['$in'].append(_ID)
                    _equip_filter = equip_filter 
                    if col_list[i] == 'pv_plant':
                        data = db.plant.find_one({'_id': ObjectId(_ID)})
                        _equip_filter['PV'] = data['name']
                        #_equip_filter['lgroup'] = {'$exists': False}
                    else:
                        data = db.equipment.find_one({'_id': ObjectId(_ID)})
                        equipment_buffer[str(data['_id'])] = data
                        if col_list[i] == 'pv_lgroup':
                            _equip_filter['PV'] = data['PV']
                            _equip_filter['lgroup'] = data['name']
                            #_equip_filter['group'] = {'$exists': False}
                        elif col_list[i] == 'pv_group':
                            _equip_filter['PV'] = data['PV']
                            _equip_filter['lgroup'] = data['lgroup']
                            _equip_filter['group'] = data['name']
                        elif col_list[i] == 'iot':
                            for iot in db.iot.find({'_id': ObjectId(_ID)}):
                                iot_buffer[str(iot['_id'])] = iot
                                alarm_filter['ID']['$in'].append(str(iot['_id']))
                        else: # inv, string, sun, wind, etc...
                            _equip_filter['_id'] = ObjectId(_ID)

                    #print(_equip_filter)
                    if col_list[i] == 'pv_group' and (_equip_filter.get('true_type', '') == 'iot' or request_dict['equip_type'] == 'all'):   #要多找iot
                        for iot in db.iot.find(_equip_filter):
                            iot_buffer[str(iot['_id'])] = iot
                            alarm_filter['ID']['$in'].append(str(iot['_id']))
                    for equip in db.equipment.find(_equip_filter):
                        equipment_buffer[str(equip['_id'])] = equip
                        alarm_filter['ID']['$in'].append(str(equip['_id']))
                except:
                    continue
        except Exception as e:
            print(e)
    # else:
    #     plant_filter = {}
    #     find_user = find_user_from_current_user()

    #     user_c=list(db.users.find({"user_id" : find_user}))[0]
    #     if user_c['plant'][0] != 'total':
    #         for i in user_c['plant']:
    #             if 'name' not in plant_filter:
    #                 plant_filter['name'] = {'$in': []}
    #             plant_filter['name']['$in'].append(i)
    #     for plant in db.plant.find(plant_filter):
    #         if request_dict['equip_type'] in ['all','pv_plant']: 
    #             alarm_filter['ID']['$in'].append(str(plant['_id']))
    #         equip_filter['PV'] = plant.get('name', '')
    #         for equip in db.equipment.find(equip_filter):
    #             equipment_buffer[str(equip['_id'])] = equip
    #             alarm_filter['ID']['$in'].append(str(equip['_id']))
    #             if equip.get('type', None) == 'pv_group' and (equip_filter.get('true_type', '') == 'iot' or request_dict['equip_type'] == 'all'):
    #                 for iot in db.iot.find({'PV': plant.get('name', ''), 'lgroup': equip['lgroup'], 'group': equip['name']}):
    #                     iot_buffer[str(iot['_id'])] = iot
    #                     alarm_filter['ID']['$in'].append(str(iot['_id']))   # iot
    if request_dict.get('alarm_group', '') != 'all':
        try:
            alarm_filter['group'] = request_dict['alarm_group']
        except Exception as e:
            print(e)
    if request_dict.get('alarm_type', '') != 'all':
        try:
            if request_dict['alarm_type'] == 'not_returned':
                alarm_filter['returntime'] = ''
            elif request_dict['alarm_type'] == 'returned':
                alarm_filter['returntime'] = {'$ne': ''}
                #alarm_filter['return_state'] = 1
            elif request_dict['alarm_type'] == 'not_archived':
                alarm_filter['show'] = 1
            elif request_dict['alarm_type'] == 'archived':
                alarm_filter['show'] = 2
        except Exception as e:
            print(e)
    #print(alarm_filter)
    alarm_data = []
    alarm_total = db.alarm.count_documents(alarm_filter)
    documents_per_page = request_dict.get('documents_per_page', 10)
    total_page = math.ceil(alarm_total/documents_per_page)
    page = request_dict.get('page', 1)
    #print(alarm_filter)
    if type(page) != int:
        print(400, 'Bad Request. Error page. Should be integer.')
    for alarm in db.alarm.find(alarm_filter).skip((page-1)*documents_per_page).limit(documents_per_page).sort('time', -1):
        try:
            place_str = ''   # 類型 廠區 分組 分區
            equip_type = ''
            equip_name = ''
            system_translate = {"DG": "DG","PV": "地面型","BESS": "屋頂型","WT": "水面型"}
            type_translate = {"inv": "變流器", "string": "串電流錶", "io": "開關", "sun": "日照計", 
            "temp": "溫度計", "wind": "風速計", "meter": "智慧電錶", "pv_meter": "智慧電錶"}
            if alarm['ID'] not in equipment_buffer and alarm['ID'] not in iot_buffer:   #代表是plant
                _data = db.plant.find_one({'_id': ObjectId(alarm['ID'])})
                place_str = system_translate.get(_data.get('type',''),_data.get('type','')) + '-' + _data.get('name', '')
                equip_type = '案場'
                equip_name = _data.get('name', '')
            elif alarm['ID'] in iot_buffer:   #代表是iot
                _data = iot_buffer[alarm['ID']]
                place_str = '-'.join([system_translate.get(plant_buffer.get(_data.get('PV')).get('type'), ''), _data.get('PV', ''), _data.get('lgroup',''), _data.get('group', '')])
                equip_type = '資料收集器'
                equip_name = _data.get('name',  '資料收集器')
            else:   # those in equipment
                _data = equipment_buffer[alarm['ID']]
                place_str = '-'.join([
                    system_translate.get(
                        plant_buffer.get(
                            _data.get('PV') if isinstance(_data.get('PV'), str) else _data.get('PV', [''])[0])
                        .get('type'), ''), 
                        _data.get('PV') if isinstance(_data.get('PV'), str) else _data.get('PV', [''])[0], 
                    ''
                ])
                equip_name = _data.get('name',  '')
                if _data.get('type') == 'pv_lgroup':
                    place_str += _data.get('name', '')
                    equip_type = '分區' if _data.get('Devide_type', '') == '' else type_translate.get(_data.get('Device_type', ''), _data.get('Device_type', ''))
                elif _data.get('type') == 'pv_group':
                    place_str += '-'.join([_data.get('lgroup', ''), _data.get('name', '')])
                    equip_type = '分組' if _data.get('Devide_type', '') == '' else type_translate.get(_data.get('Device_type', ''), _data.get('Device_type', ''))
                else:
                    place_str += '-'.join([
                        _data.get('lgroup', '') if isinstance(_data.get('lgroup', ''), str) else _data.get('lgroup', [''])[0], 
                        _data.get('group', '') if isinstance(_data.get('group', ''), str) else _data.get('group', [''])[0]
                    ])
                    equip_type = type_translate.get(_data.get('type', ''), _data.get('type', ''))
            # Get Dispatch Name
            try:
                dispatch_name = None if alarm.get('dispatch_ID', None) == None else db.dispatch.find_one({'_id': ObjectId(alarm.get('dispatch_ID', None))})['name']
            except:
                dispatch_name = None
            _dict = {
                '_id': str(alarm['_id']), 'ID': str(alarm['ID']),
                'alarm_place': place_str,
                'alarm_group': alarm.get('group', ''), 'alarm_event': alarm.get('event', ''),
                'equip_type': equip_type, 'equip_name': equip_name,
                'about': alarm.get('about', None),
                'level': alarm.get('level', 4),
                'time': datetime.strftime(alarm.get('time'), '%Y-%m-%d %H:%M:%S'),
                'checktime': '' if alarm.get('checktime', '') == '' else datetime.strftime(alarm.get('checktime'), '%Y-%m-%d %H:%M:%S'),
                'returntime': '' if alarm.get('returntime', '') == '' else datetime.strftime(alarm.get('returntime'), '%Y-%m-%d %H:%M:%S'),
                'dispatchRecord': '',
                'dispatch_ID': alarm.get('dispatch_ID', None),
                'dispatch_name': dispatch_name,
                'tools': {
                    'check': str(alarm['_id']), 'manual': str(alarm['_id']), 'archived': str(alarm['_id'])
                },
                'losing_kwh': alarm.get("losing_kwh", 0)
            }
            alarm_data.append(_dict)
        except Exception as e:
            print("mongo.py alarm_get", e)
            continue
    return {'data': alarm_data, 'total_page': total_page, 'current_page': page, 'total_count': alarm_total}

def alarm_table_data(alarms):
    try:
        alarm_option = "all"
        alarms = alarms

        return_data = {"losing":{}}
        table_dict = {}
        return_list = []
        nowtime = datetime.now()
        if(alarm_option == "設備" or alarm_option == "軟體"):
            for alarm in alarms:
                if alarm["alarm_group"] == alarm_option:
                    if alarm["alarm_event"] not in return_data:
                        return_data[alarm["alarm_event"]] = 0
                        return_data["losing"][alarm["alarm_event"]] = 0
                    return_data[alarm["alarm_event"]] += 1
                    return_data["losing"][alarm["alarm_event"]] += round(alarm.get("losing_kwh", 0), 2)


                    table_key_name = alarm.get("alarm_place", "") + "_" + alarm.get("equip_name", "") + "_" + alarm.get("alarm_event")
                    if table_key_name not in table_dict:
                        table_dict[table_key_name] = {
                            "alarm_place": alarm.get("alarm_place", "---"),
                            "equip_name" : alarm.get("equip_name", "---"),
                            "alarm_event": alarm.get("alarm_event", "---"),
                            "occur_time": 0,
                            "duration": timedelta(seconds=0)
                        }
                    
                    return_time = alarm.get("returntime", "")
                    time = alarm.get("time", nowtime)
                    duration = "---"
                    time = datetime.strptime(time, '%Y-%m-%d %H:%M:%S')
                    if len(return_time) > 0:
                        return_time = datetime.strptime(return_time, '%Y-%m-%d %H:%M:%S')
                        if return_time < time:
                            duration = time - return_time
                        else:
                            duration = return_time - time
                    else:
                        duration = nowtime - time
                    table_dict[table_key_name]["occur_time"] += 1
                    if duration != "---":
                        table_dict[table_key_name]["duration"] += duration

        elif alarm_option == "all":
            for alarm in alarms:
                if alarm["alarm_group"] == "設備" or alarm["alarm_group"] == "軟體":
                    if alarm["alarm_event"] not in return_data:
                        return_data[alarm["alarm_event"]] = 0
                        return_data["losing"][alarm["alarm_event"]] = 0
                    return_data[alarm["alarm_event"]] += 1
                    return_data["losing"][alarm["alarm_event"]] += round(alarm.get("losing_kwh", 0), 2)

                    table_key_name = alarm.get("alarm_place", "") + "_" + alarm.get("equip_name", "") + "_" + alarm.get("alarm_event")
                    if table_key_name not in table_dict:
                        table_dict[table_key_name] = {
                            "alarm_place": alarm.get("alarm_place", "---"),
                            "equip_name" : alarm.get("equip_name", "---"),
                            "alarm_event": alarm.get("alarm_event", "---"),
                            "occur_time": 0,
                            "duration": timedelta(seconds=0)
                        }
                    
                    return_time = alarm.get("returntime", "")
                    time = alarm.get("time", nowtime)
                    duration = "---"
                    time = datetime.strptime(time, '%Y-%m-%d %H:%M:%S')
                    if len(return_time) > 0:
                        return_time = datetime.strptime(return_time, '%Y-%m-%d %H:%M:%S')
                        if return_time < time:
                            duration = time - return_time
                        else:
                            duration = return_time - time
                    else:
                        duration = nowtime - time
                    table_dict[table_key_name]["occur_time"] += 1
                    if duration != "---":
                        table_dict[table_key_name]["duration"] += duration
        for data in table_dict.values():
            data["duration"] = round(data["duration"].total_seconds()/(60*60), 1)
            return_list.append(data)
        output_data = {
            "data": return_data,
            "table_data": return_list
        }
        return {'data': output_data}
    except Exception as e:
        print("mongo.py alarm_table_data", e)

def alarm_table_list(table_data):
    header = ["廠區", "設備名稱", "警報名稱", "發生次數", "持續時間(小時)"]
    table_list = []
    for data in table_data:
        row = []
        row = [data["alarm_place"], data["equip_name"], data["alarm_event"], data["occur_time"], data["duration"]]
        table_list.append(row)
    return {"header": header, "table_list": table_list}

#---------------------------派工資訊-----------------------------------------------------
def get_dispatch_finish(db, request_dict):
    try:
        dispatch_collection = db["dispatch"]
        equipment_collection = db["equipment"]
        dispatch_filter = {}
        plant = request_dict["plant"]
        if plant["ID"] != None:
            dispatch_filter["ID"] = {"$in": []}
        if plant["col"] == "pv_lgroup":
            dispatch_filter["ID"] = plant["ID"]
        elif plant["col"] == "pv_plant":
            ID_list = []
            plant_name = db["plant"].find_one({"_id": ObjectId(plant["ID"])})["name"]
            for lgroup in db["equipment"].find({"PV": plant_name, "type":"pv_lgroup"}):
                ID_list.append(str(lgroup["_id"]))
            dispatch_filter["ID"]["$in"] = ID_list
        elif plant["col"] == "pv_group":
            ID_list = []
            group = db["equipment"].find_one({"_id": ObjectId(plant["ID"])})
            for lgroup in db["equipment"].find({"PV": group["PV"], "name": group["lgroup"], "type": "pv_lgroup"}):
                ID_list.append(str(lgroup["_id"]))
            dispatch_filter["ID"]["$in"] = ID_list

        today = datetime.today()
        start_date = datetime(year=today.year, month=today.month, day=1)
        end_date = datetime(year=today.year, month=today.month+1, day=1) - timedelta(days=1)
        dispatch_filter['dispatch_time'] = {'$gte': start_date, '$lt': end_date}

        type_dict = {
            "alarm": "告警檢修",
            "wash": "清洗",
            "regular": "定檢"
        }
        stage_dict = {
            "wait_for_priority": "等待優先度排序",
            "wait_for_take": "待接單",
            "merged": "已合併",
            "took_wait_date_enter": "等待輸入派工日期",
            "wait_for_dispatch": "等待派工",
            "wait_admin_confirm_date": "等待管理人員確認派工日期",
            "dispatched_wait_for_review": "等待AI驗收",
            "auto_reviewed_wait_for_manual": "等待管理人員協助驗收",
            "review_failed": "驗收失敗",
            "dispatch_finish": "已完成工單"
        }
        _dispatch_data = []
        for dispatch_data in dispatch_collection.find(dispatch_filter):
            #廠區
            ID = dispatch_data.get("ID", "")
            equipment = equipment_collection.find_one({"_id": ObjectId(ID), "type": "pv_lgroup"})
            station = equipment.get("PV", "") + "/" + equipment.get("name", "")
            #種類
            type_name = ""
            for index, types in enumerate(dispatch_data.get("type", [])):
                if types in type_dict:
                    if index > 0:
                        type_name += "、"
                    type_name += type_dict[types]
            #狀態
            _stage = dispatch_data.get("stage", "")
            stage = stage_dict[_stage]
            #維運人員
            maintainer = "---"
            maintainer_ID = dispatch_data.get("maintainer_ID", "")
            user = db["users"].find_one({"_id": ObjectId(maintainer_ID)})
            if user:
                user_data = user.get("user_data", {})
                if user_data != {}:
                    _name = user_data.get("name", "")
                    if len(_name) > 0:
                        maintainer = _name
            #派工日期
            _time = dispatch_data.get("dispatch_time", "")
            if _time != "" and _time != None:
                dispatch_time = datetime.strftime(_time, '%Y-%m-%d')
            
            _dispatch_data.append([station, type_name, stage, maintainer, dispatch_time])
            
        return _dispatch_data
    except Exception as e:
        print("mongo.py get_dispatch_finish",e)

def get_dispatch_init(db, request_dict):
    try:
        dispatch_collection = db["dispatch"]
        equipment_collection = db["equipment"]
        dispatch_filter = {}
        plant = request_dict["plant"]
        if plant["ID"] != None:
            dispatch_filter["ID"] = {"$in": []}
        if plant["col"] == "pv_lgroup":
            dispatch_filter["ID"] = plant["ID"]
        elif plant["col"] == "pv_plant":
            ID_list = []
            plant_name = db["plant"].find_one({"_id": ObjectId(plant["ID"])})["name"]
            for lgroup in db["equipment"].find({"PV": plant_name, "type":"pv_lgroup"}):
                ID_list.append(str(lgroup["_id"]))
            dispatch_filter["ID"]["$in"] = ID_list
        elif plant["col"] == "pv_group":
            ID_list = []
            group = db["equipment"].find_one({"_id": ObjectId(plant["ID"])})
            for lgroup in db["equipment"].find({"PV": group["PV"], "name": group["lgroup"], "type": "pv_lgroup"}):
                ID_list.append(str(lgroup["_id"]))
            dispatch_filter["ID"]["$in"] = ID_list

        today = datetime.today()
        start_date = datetime(year=today.year, month=today.month, day=1)
        end_date = datetime(year=today.year, month=today.month+1, day=1) - timedelta(days=1)
        dispatch_filter['init_time'] = {'$gte': start_date, '$lt': end_date}

        type_dict = {
            "alarm": "告警檢修",
            "wash": "清洗",
            "regular": "定檢"
        }
        stage_dict = {
            "wait_for_priority": "等待優先度排序",
            "wait_for_take": "待接單",
            "merged": "已合併",
            "took_wait_date_enter": "等待輸入派工日期",
            "wait_for_dispatch": "等待派工",
            "wait_admin_confirm_date": "等待管理人員確認派工日期",
            "dispatched_wait_for_review": "等待AI驗收",
            "auto_reviewed_wait_for_manual": "等待管理人員協助驗收",
            "review_failed": "驗收失敗",
            "dispatch_finish": "已完成工單"
        }
        _dispatch_data = []
        for dispatch_data in dispatch_collection.find(dispatch_filter):
            if dispatch_data.get("dispatch_time", "") == "" or dispatch_data.get("dispatch_time", "") == None:
                #廠區
                ID = dispatch_data.get("ID", "")
                equipment = equipment_collection.find_one({"_id": ObjectId(ID), "type": "pv_lgroup"})
                station = equipment.get("PV", "") + "/" + equipment.get("name", "")
                #種類
                type_name = ""
                for index, types in enumerate(dispatch_data.get("type", [])):
                    if types in type_dict:
                        if index > 0:
                            type_name += "、"
                        type_name += type_dict[types]
                #狀態
                _stage = dispatch_data.get("stage", "")
                stage = stage_dict[_stage]
                #維運人員
                maintainer = "---"
                maintainer_ID = dispatch_data.get("maintainer_ID", "")
                user = db["users"].find_one({"_id": ObjectId(maintainer_ID)})
                if user:
                    user_data = user.get("user_data", {})
                    if user_data != {}:
                        _name = user_data.get("name", "")
                        if len(_name) > 0:
                            maintainer = _name
                #派工日期
                _time = dispatch_data.get("init_time", "")
                if _time != "" and _time != None:
                    dispatch_time = datetime.strftime(_time, '%Y-%m-%d')
                
                _dispatch_data.append([station, type_name, stage, maintainer, dispatch_time])
            
        return _dispatch_data
    except Exception as e:
        print("mongo.py get_dispatch_init",e)
#---------------------------報表製作-----------------------------------------------------

#初始設定
def document_initial(doc):
    #設定邊界
    section = doc.sections[0]
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    #設定字型
    doc.styles['Normal'].font.size = Pt(12)
    #設定字體(英文)
    doc.styles['Normal'].font.name = 'Times New Roman'
    #中文
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

#資料表格
def add_table(doc, head_datas, datas, width, style=None):
    table = doc.add_table(rows=1, cols=len(head_datas), style=("Table Grid" if style is None else style))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    head_cells = table.rows[0].cells
    #表格標頭文字
    for index, head_item in enumerate(head_datas):
        head_cells[index].text = head_item
        head_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        head_cells[index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #表格資料
    for data in datas:
        row_cells = table.add_row().cells
        for index, cell in enumerate(row_cells):
            cell.text = str(data[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # #表格內寬度
    # for index, column in enumerate(table.columns):
    #     for cell in column.cells:
    #         cell.width = Cm(width[index])

    #表格內高度
    for index, row in enumerate(table.rows):
        if index == 0:
            row.height = Cm(0.6)
        else:
            row.height = Cm(0.64)

#刪除段落、圖片等
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    #p._p = p._element = None
    paragraph._p = paragraph._element = None

#資訊欄
def imformation_table(doc, imformation_data):
    width = [3.9, 15.1]
    rows = 0
    index = 0
    keys = []
    row_index = 0
    key_index = 0

    #設定標題
    heading = doc.add_paragraph("")
    heading.style = doc.styles['Heading 1']
    #段前間距
    heading.paragraph_format.space_before = Pt(0)
    #段後間距
    heading.paragraph_format.space_after = Pt(15)

    run = heading.add_run("太陽光電系統維運每週報表")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

    
    # run.space_before = Pt(9)
    # run.space_after = Pt(9)
    # run.line_spacing = 1.5

    for value in imformation_data.values():
        if type(value) == list:
            rows += len(value)
        elif type(value) == str:
            rows += 1
    
    table = doc.add_table(rows=rows, cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    #table.cell(0, 0).text = "表角"

    for key in imformation_data.keys():
        keys.append(key)

    while row_index < rows:
        if type(imformation_data[keys[key_index]]) == list:
            for content in imformation_data[keys[key_index]]:
                row = table.rows[row_index]
                row.cells[0].text = keys[key_index]
                row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row.cells[1].text = content
                row_index += 1
        else:
            row = table.rows[row_index]
            row.cells[0].text = keys[key_index]
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row.cells[1].text = imformation_data[keys[key_index]]
            row_index += 1
        key_index += 1

    for index, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = Cm(width[index])


    #合併重複名稱
    col = table.columns[0]
    row_index = 0
    while row_index < rows-1:
        if col.cells[row_index].text == col.cells[row_index+1].text:
            content = col.cells[row_index].text
            col.cells[row_index].merge(col.cells[row_index+1])
            col.cells[row_index].text = content
            col.cells[row_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            col.cells[row_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_index += 1

#設定頁首
def header(doc, imformations, logo=None):
    section = doc.sections[0]
    header = section.header
    # paragraph = header.paragraphs[0]
    paragraph = header.add_paragraph()
    #段落間距
    paragraph.paragraph_format.space_after = Pt(0)
    #單行間距(長度為絕對距離，浮點數為間距)
    paragraph.paragraph_format.line_spacing = 1.0
    #刪除首段
    delete_paragraph(header.paragraphs[0])
    tab_stops = paragraph.paragraph_format.tab_stops
    logo_run = paragraph.add_run()

    if logo:
        logo_run.add_picture(logo, height=Inches(0.65))
    text_run = paragraph.add_run('\t')
    text_run = paragraph.add_run(imformations[0])
    #設置右邊文字位置tab
    margin_end = Inches(section.page_width.inches - (section.left_margin.inches + section.right_margin.inches))
    tab_stop = tab_stops.add_tab_stop(margin_end, WD_TAB_ALIGNMENT.RIGHT)
    #字體大小、斜體
    text_run.font.size = Pt(14)
    text_run.font.italic = True

    for index, imformation in enumerate(imformations):
        if index == 0:
            #paragraph = header.paragraphs[0]
            continue
        else:
            paragraph = header.add_paragraph()
            #pass
        run = paragraph.add_run(""+imformation)
        if index == 0:
            #run.font.size = Pt(14)
            continue
        else:
            run.font.size = Pt(10)
        run.font.italic = True
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
def table_data(dates, irrh_data, meter_data, pr_data):
    data = []
    for date, irrh, meter, pr in zip(dates, irrh_data, meter_data, pr_data):
        data.append([str(date).replace("-", "/"), irrh, meter, pr])
    return data

def report(pv, name, header_imformation, imformation, table_head_datas, table_datas, ID, time_interval, logo_path, path=None, alarm_data=None, dispatch_data=None):
    logo = logo_path

    #總長約19cm
    width = [4.71, 5.49, 4.4, 4.4]

    doc = Document()
    #初始設定
    document_initial(doc)
    #設定頁首
    try:
        header(doc, header_imformation, logo)
    except:
        header(doc, header_imformation)
    #資訊欄
    imformation_table(doc, imformation)
    p = doc.add_paragraph()
    #資料表格
    add_table(doc, table_head_datas, table_datas, width=width)

    if alarm_data != None:
        p = doc.add_paragraph()
        heading = doc.add_paragraph("")
        heading.style = doc.styles['Heading 1']
        #段前間距
        heading.paragraph_format.space_before = Pt(0)
        #段後間距
        heading.paragraph_format.space_after = Pt(15)

        run = heading.add_run("警報列表")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

        #警報列表
        if alarm_data != None and alarm_data.get("table_list", []) != []:
            add_table(doc, alarm_data["header"], alarm_data["table_list"], width=width)
        else:
            add_table(doc, ["無警報"], [], width=width)

        #派工列表
        dispatch_finish_header = ["廠區", "類別", "狀態", "維運人員", "派工日期"]
        dispatch_init_header = ["廠區", "類別", "狀態", "維運人員", "工單產生日期"]

        p = doc.add_paragraph()
        heading = doc.add_paragraph("")
        heading.style = doc.styles['Heading 1']
        #段前間距
        heading.paragraph_format.space_before = Pt(0)
        #段後間距
        heading.paragraph_format.space_after = Pt(15)

        run = heading.add_run("派工列表")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
        if dispatch_data != None:
            if dispatch_data[0] != []:
                add_table(doc, dispatch_finish_header, dispatch_data[0], width)
            else:
                add_table(doc, ["無派工"], [], width=width)

        p = doc.add_paragraph()
        heading = doc.add_paragraph("")
        heading.style = doc.styles['Heading 1']
        #段前間距
        heading.paragraph_format.space_before = Pt(0)
        #段後間距
        heading.paragraph_format.space_after = Pt(15)

        run = heading.add_run("派工列表-未派工")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
        if dispatch_data != None:
            if dispatch_data[1] != []:
                add_table(doc, dispatch_init_header, dispatch_data[1], width)
            else:
                add_table(doc, ["無派工"], [], width=width)

    today = date.today()
    delta = timedelta(days=1)
    yesterday = today - delta
    year = yesterday.year
    month = yesterday.month

    #將檔案放至資料夾
    # if time_interval == "day":
    #     if month < 10:
    #         doc_name = name + "_" + str(year) + "-" + "0" + str(month) + ".docx"
    #     else:
    #          doc_name = name + "_" + str(year) + "-" + str(month) + ".docx"
    # elif time_interval == "month":
    #     doc_name = name + "_" + str(year) + ".docx"
    # else:
    #     doc_name = name + "_" + str(yesterday) + "_" + time_interval + ".docx"
    if path:
        file_position = path+"/ID_interval_word"
    else:
        file_position = os.getcwd()+"/ID_interval_word"
    if not os.path.isdir(file_position):
        os.mkdir(file_position)
    file_position += "/"+ID
    if not os.path.isdir(file_position):
        os.mkdir(file_position)
    file_position += "/"+time_interval
    if not os.path.isdir(file_position):
        os.mkdir(file_position)
    filename = file_position+"/"+name
    doc.save(filename)

    #將紀錄寫至mongodb
    mongo_collection = pv["ID_interval_word"]
    record_dict = {
        "ID": ID,
        "time_interval": time_interval,
        "filename": name,
        "show": 1
    }
    check_dict = mongo_collection.find_one(record_dict)
    if check_dict:
        mongo_collection.delete_one(check_dict)
    mongo_collection.insert_one(record_dict)

if __name__ == "__main__":
    pv = mongo_connect()
    plant_id = "5e8d4c884a11d7e11cd2050e"
    equipment_id = "5e8d4c884a11d7e11cd20521"
    test_id = "5e12d4f67dda745b3bd503bf"
    plant_id, solar_ID, meter_ID, pr_ID, collection = id_identify(pv, test_id)
    name, field_position, capacity = field_imformation(pv, plant_id)

    start_time = "2022-03-01"
    start_time = datetime.strptime(start_time, '%Y-%m-%d')
    end_time = "2022-03-02"
    end_time = datetime.strptime(end_time, '%Y-%m-%d')
    time_interval = "hour"
    time_list = set_time_interval(start_time, end_time, time_interval)

    irrh_data = irrh_cal(pv, solar_ID, time_list, time_interval)
    meter_data = meter_cal(pv, meter_ID, time_list, time_interval)
    pr_data = pr_cal(pv, pr_ID, time_list, time_interval)
    max_val = avg_max_value(pv, pr_ID, collection, start_time, end_time, avg_status=False)


    header_data = company_imformation(pv, plant_id)

    name = project_name(pv, pr_ID, time_interval)
    date = str(start_time) + "~" + str(end_time)
    imformation_dict = imformation_data(name, date, field_position, str(capacity), str(max_val))
    table_head_datas = ["時間", "日照量", "發電量", "PR"]
    data = table_data(time_list, irrh_data, meter_data, pr_data)
    report(pv, name, header_data, imformation_dict, table_head_datas, data, test_id, time_interval, "/images/logo.jpg")
    