from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.shared import RGBColor

#初始設定
def document_initial(doc):
    #設定邊界
    section = doc.sections[0]
    section.left_margin = Cm(1.27)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)
    #設定字型
    doc.styles['Normal'].font.size = Pt(12)
    #設定字體(英文)
    doc.styles['Normal'].font.name = 'Times New Roman'
    #中文
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

#資料表格
def add_table(doc, head_datas, datas, width, style=None):
    table = doc.add_table(rows=1, cols=len(head_datas), style=("Table Grid" if style is None else style))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    head_cells = table.rows[0].cells
    #表格標頭文字
    for index, head_item in enumerate(head_datas):
        head_cells[index].text = head_item
        head_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        head_cells[index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #表格資料
    for data in datas:
        row_cells = table.add_row().cells
        for index, cell in enumerate(row_cells):
            cell.text = str(data[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #表格內高度
    for index, row in enumerate(table.rows):
        if index == 0:
            row.height = Cm(0.6)
        else:
            row.height = Cm(0.64)

    #表格內寬度
    for index, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = Cm(width[index])

#刪除段落、圖片等
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    #p._p = p._element = None
    paragraph._p = paragraph._element = None

#資訊欄
def imformation_table(doc, imformation_data):
    width = [3.9, 15.1]
    rows = 0
    index = 0
    keys = []
    row_index = 0
    key_index = 0


    #設定標題
    heading = doc.add_paragraph("")
    heading.style = doc.styles['Heading 1']
    #段前間距
    heading.paragraph_format.space_before = Pt(0)
    #段後間距
    heading.paragraph_format.space_after = Pt(15)

    run = heading.add_run("太陽光電系統維運每週報表")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

    
    # run.space_before = Pt(9)
    # run.space_after = Pt(9)
    # run.line_spacing = 1.5

    for value in imformation_data.values():
        if type(value) == list:
            rows += len(value)
        elif type(value) == str:
            rows += 1
    
    table = doc.add_table(rows=rows, cols=2, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    #table.cell(0, 0).text = "表角"

    for key in imformation_data.keys():
        keys.append(key)

    while row_index < rows:
        if type(imformation_data[keys[key_index]]) == list:
            for content in imformation_data[keys[key_index]]:
                row = table.rows[row_index]
                row.cells[0].text = keys[key_index]
                row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row.cells[1].text = content
                row_index += 1
        else:
            row = table.rows[row_index]
            row.cells[0].text = keys[key_index]
            row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row.cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row.cells[1].text = imformation_data[keys[key_index]]
            row_index += 1
        key_index += 1

    for index, column in enumerate(table.columns):
        for cell in column.cells:
            cell.width = Cm(width[index])


    #合併重複名稱
    col = table.columns[0]
    row_index = 0
    while row_index < rows-1:
        if col.cells[row_index].text == col.cells[row_index+1].text:
            content = col.cells[row_index].text
            col.cells[row_index].merge(col.cells[row_index+1])
            col.cells[row_index].text = content
            col.cells[row_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            col.cells[row_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        row_index += 1

#設定頁首
def header(doc, imformations, logo):
    section = doc.sections[0]
    header = section.header
    # paragraph = header.paragraphs[0]
    paragraph = header.add_paragraph()
    #段落間距
    paragraph.paragraph_format.space_after = Pt(0)
    #單行間距(長度為絕對距離，浮點數為間距)
    paragraph.paragraph_format.line_spacing = 1.0
    #刪除首段
    delete_paragraph(header.paragraphs[0])
    tab_stops = paragraph.paragraph_format.tab_stops
    logo_run = paragraph.add_run()


    logo_run.add_picture(logo, height=Inches(0.65))
    text_run = paragraph.add_run('\t')
    text_run = paragraph.add_run(imformations[0])
    #設置右邊文字位置tab
    margin_end = Inches(section.page_width.inches - (section.left_margin.inches + section.right_margin.inches))
    tab_stop = tab_stops.add_tab_stop(margin_end, WD_TAB_ALIGNMENT.RIGHT)
    #字體大小、斜體
    text_run.font.size = Pt(14)
    text_run.font.italic = True

    for index, imformation in enumerate(imformations):
        if index == 0:
            #paragraph = header.paragraphs[0]
            continue
        else:
            paragraph = header.add_paragraph()
            #pass
        run = paragraph.add_run(""+imformation)
        if index == 0:
            #run.font.size = Pt(14)
            continue
        else:
            run.font.size = Pt(10)
        run.font.italic = True
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    
 
def main():
    logo = "title_picture.jpg"
    header_imformation = ["全勤能源科技服務股份有限公司", "台北市松山區八德路二段412號2樓Tel: (02)8771-0081"]

    imformation = {
        "專案名稱": "上萬安段-上萬安段", 
        "日期": "2022-02-07 00:00:00~2022-02-14 00:00:00",
        "案場位置": ["地址 : 屏東縣新埤鄉永新路6-1號", "經緯度 : 22°30'41.2\"N 120°34'25.9\"E"],
        "設置容量": "1884.18 kWp"
    }
    
    #總長約19cm
    width = [4.71, 5.49, 4.4, 4.4]
    table_head_datas = ["週次", "週平均DMY(kwh/kwp/day)", "週總發電量(kwh)", "週PR(%)"]
    table_datas = [
        ["2021/12/27", "11.89", "20401.8kWh", "87.75%"],
        ["2022/1/24", "2.78", "36691.2kWh", "86.89%"]
    ]

    doc = Document()
    #初始設定
    document_initial(doc)
    #設定頁首
    header(doc, header_imformation, logo)
    #資訊欄
    imformation_table(doc, imformation)
    p = doc.add_paragraph()
    #資料表格
    add_table(doc, table_head_datas, table_datas, width=width)
    doc.save("周報表.docx")

if __name__ == "__main__":
    main()