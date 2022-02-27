from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

#初始設定
def document_initial(doc):
    #設定邊界
    section = doc.sections[0]
    #設定字型大小
    doc.styles['Normal'].font.size = Pt(14)
    #設定字體(英文)
    doc.styles['Normal'].font.name = 'Times New Roman'
    #中文
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')

#資料表格
def add_table(doc, datas, style=None, merge_data=None):
    head_datas = datas.pop(0)
    table = doc.add_table(rows=1, cols=len(head_datas), style=("Table Grid" if style is None else style))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    head_cells = table.rows[0].cells
    #表格標頭文字
    for index, head_item in enumerate(head_datas):
        head_cells[index].text = head_item
        head_cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        head_cells[index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #表格資料
    for data in datas:
        row_cells = table.add_row().cells
        for index, cell in enumerate(row_cells):
            cell.text = str(data[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #表格內高度
    for index, row in enumerate(table.rows):
        if index == 0:
            row.height = Cm(0.6)
        else:
            row.height = Cm(0.64)

    if merge_data:
        for data in merge_data:
            start_row = data[0]
            start_column = data[1]
            end_row = data[2]
            end_column = data[3]
            content = table.cell(start_row, start_column).text
            table.cell(start_row, start_column).merge(table.cell(end_row, end_column))
            table.cell(start_row, start_column).text = content
            table.cell(start_row, start_column).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            table.cell(start_row, start_column).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#刪除段落、圖片等
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    #p._p = p._element = None
    paragraph._p = paragraph._element = None

def main():
    table_datas_1 = [
        ["U1", "Mode", "Rated(kV)", "%V", "Vangle", "3φ MVAsc", "Connection", "X/R"],
        ["U1", "Swing", "100", "100", "0", "100", "Y solid", "35"]
    ]

    merge_data = [[0, 0, 1, 0]]

    table_datas_2 = [
        ["Impedance Number", "R", "X"],
        ["Z1", "2", "4"],
        ["Z2", "1", "3"],
        ["Z3", "1.25", "2.5"]
    ]

    table_datas_3 = [
        ["Load Number", "P (MW)", "Q (Mvar)", "Connection"],
        ["Load1", "256.6", "110.2", "Y solid"],
        ["Load2", "138.6", "45.2", "Y solid"]
    ]

    table_datas_4 = [
        ["Bus Number", "Rated(kV)", "%V", "Vangle"],
        ["Bus1", "100", "100", "0"],
        ["Bus2", "100", "100", "0"],
        ["Bus3", "100", "100", "0"]
    ]

    doc = Document()
    #初始設定
    document_initial(doc)

    #資料表格
    add_table(doc, table_datas_1, merge_data=merge_data)
    #新增段落(一行)
    doc.add_paragraph()

    add_table(doc, table_datas_4)

    doc.add_paragraph()

    add_table(doc, table_datas_2)

    doc.add_paragraph()

    add_table(doc, table_datas_3)


    doc.save("homework.docx")

if __name__ == "__main__":
    main()